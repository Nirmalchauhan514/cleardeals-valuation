import streamlit as st
import pandas as pd
from docx import Document
from datetime import date

st.set_page_config(page_title="ClearDeals Valuation Tool", layout="centered")

st.title("🏡 ClearDeals - Gandhinagar Property Valuation Tool")
st.write("Generate instant property valuation report based on current market prices in Gandhinagar.")

# --- Area & Property Price Map (Customizable as per market)
price_data = {
    "Kudasan": {"1 BHK Flat": 3900, "2 BHK Flat": 4200, "3 BHK Flat": 4500, "Villa": 6000, "Shop": 11000},
    "Sargasan": {"1 BHK Flat": 3800, "2 BHK Flat": 4000, "3 BHK Flat": 4300, "Villa": 5800, "Shop": 10000},
    "Raysan": {"1 BHK Flat": 3600, "2 BHK Flat": 3900, "3 BHK Flat": 4200, "Villa": 5500, "Shop": 9500},
    "PDPU Road": {"2 BHK Flat": 4000, "3 BHK Flat": 4300, "Villa": 5900, "Shop": 9800},
    "Zundal": {"2 BHK Flat": 3900, "3 BHK Flat": 4200, "Villa": 5600, "Shop": 9200},
}

areas = list(price_data.keys())

# --- Form
with st.form("valuation_form"):
    area = st.selectbox("Select Area", areas)
    property_type = st.selectbox("Select Property Type", list(price_data[area].keys()))
    size = st.number_input("Enter Property Size (sq.ft)", min_value=200, step=50)
    owner_name = st.text_input("Owner's Name")
    mobile = st.text_input("Contact Number")
    submitted = st.form_submit_button("Generate Valuation Report")

# --- Valuation Logic
if submitted:
    rate = price_data[area][property_type]
    est_price = size * rate

    st.success(f"✅ Estimated Value: ₹ {est_price:,.0f}")
    st.write("You can now download your valuation report.")

    # --- Generate Word Report
    doc = Document()
    doc.add_heading('ClearDeals Property Valuation Report', 0)
    doc.add_paragraph(f"📅 Date: {date.today().strftime('%d-%m-%Y')}")
    doc.add_paragraph(f"👤 Owner Name: {owner_name}")
    doc.add_paragraph(f"📞 Contact: {mobile}")
    doc.add_paragraph(f"📍 Area: {area}")
    doc.add_paragraph(f"🏠 Property Type: {property_type}")
    doc.add_paragraph(f"📐 Size: {size} sq.ft")
    doc.add_paragraph(f"💰 Market Rate: ₹{rate}/sq.ft")
    doc.add_paragraph(f"🔎 Estimated Valuation: ₹ {est_price:,.0f}")
    doc.add_paragraph("\nThank you for using ClearDeals Gandhinagar Tool.")

    filename = f"valuation_{owner_name.replace(' ', '_')}.docx"
    doc.save(filename)

    with open(filename, "rb") as f:
        st.download_button("📥 Download Word Report", f, file_name=filename)

